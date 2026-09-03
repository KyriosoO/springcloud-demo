from __future__ import annotations

import hashlib
import io
import json
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import pymupdf
import httpx
import pytest
from pydantic import ValidationError
from docx import Document
from openpyxl import Workbook

from knowledge_corpus_tools.acquire import OfficialAssetAcquirer
from knowledge_corpus_tools.audit import _source_status
from knowledge_corpus_tools.chunking import chunk_document
from knowledge_corpus_tools.errors import ContractError, SafetyError, StateConflict
from knowledge_corpus_tools.indexing import EmbeddingClient, ElasticsearchCandidateAdmin, to_index_document
from knowledge_corpus_tools.cli import main as cli_main
from knowledge_corpus_tools.jsonio import canonical_bytes, exclusive_write, sha256_bytes, strict_loads, write_jsonl
from knowledge_corpus_tools.models import (
    CorpusSource,
    CorpusSourceCatalog,
    AssetManifest,
    AuditItem,
    AuditSummary,
    CorpusChunk,
    IntegrityFinding,
    IntegrityStatus,
    OcrStatus,
    Priority,
    ProcessingResult,
    SourceProbe,
    SourcePolicy,
    SourceStatus,
    StageAUatCase,
    StageAUatResult,
)
from knowledge_corpus_tools.parsing import discover_attachments, parse_asset
from knowledge_corpus_tools.pipeline import acquire_catalog, process_assets
from knowledge_corpus_tools.release import AliasReleaseManager
from knowledge_corpus_tools.safety import normalize_official_url, validate_zip_container

SHA = "0" * 64
HOSTS = frozenset({"www.chinatax.gov.cn"})


def test_strict_json_and_exclusive_write(tmp_path: Path) -> None:
    with pytest.raises(ContractError, match="duplicate"):
        strict_loads('{"a":1,"a":2}')
    with pytest.raises(ContractError, match="NFC"):
        strict_loads(json.dumps({"value": "e\u0301"}))
    target = tmp_path / "one.json"
    exclusive_write(target, b"one")
    with pytest.raises(StateConflict):
        exclusive_write(target, b"two")
    assert target.read_bytes() == b"one"


def test_audit_distinguishes_missing_url_from_unreachable_source() -> None:
    assert _source_status({"officialSourceUrl": ""}) is SourceStatus.URL_MISSING
    assert _source_status({"officialSourceUrl": "https://www.chinatax.gov.cn/x", "sourceHttpStatus": 403}) is SourceStatus.UNREACHABLE


def test_audit_validator_rejects_same_count_with_different_content(tmp_path: Path) -> None:
    items = tmp_path / "audit.jsonl"
    item = AuditItem(
        schema_version=2,
        document_id="doc-1",
        title="标题",
        validity_status="UNKNOWN",
        logical_domain_id="tax.policy",
        priority=Priority.P2,
        priority_reason="inventory_only",
        inventory=None,
        source=SourceProbe(status=SourceStatus.URL_MISSING, redirect_count=0),
        integrity=IntegrityFinding(status=IntegrityStatus.NOT_ASSESSABLE, reason="source_not_readable"),
        requires_human_review=True,
    )
    write_jsonl(items, (item,))
    summary = AuditSummary(
        schema_version=2,
        generated_at_utc=datetime.now(UTC),
        current_alias="read",
        current_index="index",
        current_index_uuid="uuid",
        current_document_count=0,
        current_chunk_count=0,
        audit_item_count=1,
        priority_counts={Priority.P2: 1},
        source_status_counts={SourceStatus.URL_MISSING: 1},
        integrity_status_counts={IntegrityStatus.NOT_ASSESSABLE: 1},
        es_read_requests=0,
        source_get_budget=0,
        source_get_used=0,
        retry_count=0,
        index_write_count=0,
        audit_jsonl_sha256="0" * 64,
    )
    summary_path = tmp_path / "summary.json"
    summary_path.write_bytes(canonical_bytes(summary.model_dump(mode="json")))

    with pytest.raises(SystemExit, match="hash differs"):
        cli_main(["validate-audit", "--items", str(items), "--summary", str(summary_path)])


@pytest.mark.parametrize(
    "url",
    [
        "http://www.chinatax.gov.cn/a.pdf",
        "https://127.0.0.1/a.pdf",
        "https://evil.example/a.pdf",
        "https://user:pass@www.chinatax.gov.cn/a.pdf",
        "https://www.chinatax.gov.cn:8443/a.pdf",
    ],
)
def test_official_url_rejects_unsafe_targets(url: str) -> None:
    with pytest.raises(SafetyError):
        normalize_official_url(url, HOSTS)


def test_zip_traversal_and_bomb_fail_closed() -> None:
    raw = io.BytesIO()
    with zipfile.ZipFile(raw, "w") as archive:
        archive.writestr("../escape.xml", "x")
    with zipfile.ZipFile(io.BytesIO(raw.getvalue())) as archive, pytest.raises(SafetyError, match="traversal"):
        validate_zip_container(archive)

    raw = io.BytesIO()
    with zipfile.ZipFile(raw, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "0" * 1_000_000)
    with zipfile.ZipFile(io.BytesIO(raw.getvalue())) as archive, pytest.raises(SafetyError, match="ratio"):
        validate_zip_container(archive, max_ratio=2.0)


def test_attachment_discovery_is_official_and_stable() -> None:
    raw = b'<html><body><a href="https://evil.example/page.html">ordinary link</a><a href="files/a.docx">A</a><a href="files/a.docx">again</a><a href="https://evil.example/x.pdf">bad</a></body></html>'
    with pytest.raises(SafetyError, match="allowlisted"):
        discover_attachments(
            raw,
            source_page_url="https://www.chinatax.gov.cn/root/page.html",
            parent_document_id="tax-1",
            allowed_hosts=HOSTS,
        )
    safe_raw = b'<html><body><a href="https://evil.example/page.html">ordinary link</a><a href="files/a.docx">A</a><a href="files/a.docx">again</a></body></html>'
    result = discover_attachments(
        safe_raw,
        source_page_url="https://www.chinatax.gov.cn/root/page.html",
        parent_document_id="tax-1",
        allowed_hosts=HOSTS,
    )
    assert len(result) == 1
    assert result[0].attachment_url == "https://www.chinatax.gov.cn/root/files/a.docx"


def test_acquirer_validates_redirect_mime_and_is_idempotent(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("正文")
    stream = io.BytesIO()
    document.save(stream)
    raw = stream.getvalue()

    def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path == "/start":
            return httpx.Response(302, headers={"Location": "/asset.docx"})
        return httpx.Response(
            200,
            headers={"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
            content=raw,
        )

    with httpx.Client(transport=httpx.MockTransport(handler), follow_redirects=False) as client:
        acquirer = OfficialAssetAcquirer(workspace=tmp_path, allowed_hosts=HOSTS, client=client)
        first = acquirer.fetch(
            url="https://www.chinatax.gov.cn/start",
            parent_document_id="tax-1",
            official_source_proof="official page link",
            filename="asset.docx",
        )
        second = acquirer.fetch(
            url="https://www.chinatax.gov.cn/start",
            parent_document_id="tax-1",
            official_source_proof="official page link",
            filename="asset.docx",
        )
    assert first.sha256 == second.sha256 == hashlib.sha256(raw).hexdigest()
    assert (tmp_path / first.storage_relative_path).read_bytes() == raw


def test_acquirer_records_official_docx_that_is_actually_legacy_doc(tmp_path: Path) -> None:
    raw = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\0" * 64

    def handler(_: httpx.Request) -> httpx.Response:
        return httpx.Response(
            200,
            headers={"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
            content=raw,
        )

    with httpx.Client(transport=httpx.MockTransport(handler)) as client:
        manifest = OfficialAssetAcquirer(workspace=tmp_path, allowed_hosts=HOSTS, client=client).fetch(
            url="https://www.chinatax.gov.cn/a.docx",
            parent_document_id="tax-1",
            official_source_proof="official page link",
            filename="a.docx",
        )
    assert manifest.source_extension == ".docx"
    assert manifest.extension == ".doc"
    assert manifest.format_mismatch is True


def test_acquisition_isolates_network_timeout_as_limited_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    catalog = CorpusSourceCatalog(
        schema_version=1,
        priority=Priority.P0,
        assets=(
            CorpusSource(
                document_id="tax-1",
                document_number="doc-1",
                title="标题",
                source_url="https://www.chinatax.gov.cn/source.html",
                official_replacement_proof="official page",
            ),
        ),
    )
    policy = SourcePolicy(
        schema_version=1,
        allowed_hosts=("www.chinatax.gov.cn",),
        max_asset_bytes=1024,
        max_redirects=3,
    )
    catalog_path = tmp_path / "catalog.json"
    policy_path = tmp_path / "policy.json"
    catalog_path.write_bytes(canonical_bytes(catalog.model_dump(mode="json")))
    policy_path.write_bytes(canonical_bytes(policy.model_dump(mode="json")))

    def timeout_handler(request: httpx.Request) -> httpx.Response:
        raise httpx.ReadTimeout("timed out", request=request)

    client = httpx.Client(transport=httpx.MockTransport(timeout_handler))
    monkeypatch.setattr(
        "knowledge_corpus_tools.pipeline.OfficialAssetAcquirer",
        lambda **kwargs: OfficialAssetAcquirer(client=client, **kwargs),
    )

    result = acquire_catalog(
        catalog_path=catalog_path,
        policy_path=policy_path,
        workspace=tmp_path,
        run_id="timeout-run",
        source_get_budget=1,
    )

    client.close()
    assert result.source_get_used == 1
    assert result.downloaded_asset_count == 0
    assert [failure.reason for failure in result.failures] == ["source_unreachable"]


def test_processing_isolates_malformed_office_container(tmp_path: Path) -> None:
    raw = b"PK\x03\x04not-a-valid-office-container"
    digest = hashlib.sha256(raw).hexdigest()
    relative = Path("raw") / "sha256" / digest[:2] / f"{digest}.docx"
    asset_path = tmp_path / relative
    asset_path.parent.mkdir(parents=True)
    asset_path.write_bytes(raw)
    manifest = AssetManifest(
        schema_version=1,
        asset_id=f"asset-{digest[:24]}",
        asset_version=digest,
        parent_document_id="tax-1",
        source_url="https://www.chinatax.gov.cn/a.docx",
        source_final_url="https://www.chinatax.gov.cn/a.docx",
        fetched_at_utc=datetime.now(UTC),
        filename="a.docx",
        source_extension=".docx",
        extension=".docx",
        format_mismatch=False,
        declared_mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        detected_mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        sha256=digest,
        byte_count=len(raw),
        storage_relative_path=relative.as_posix(),
        official_source_proof="official page",
    )
    manifest_path = tmp_path / "manifest.jsonl"
    write_jsonl(manifest_path, (manifest,))

    result = process_assets(
        asset_manifest_path=manifest_path,
        workspace=tmp_path,
        run_id="malformed-run",
        enable_ocr=False,
    )

    assert result.rejected_asset_count == 1
    assert result.chunk_count == 0
    assert [failure.reason for failure in result.failures] == ["parse_failed"]


def test_parsers_preserve_docx_and_xlsx_tables(tmp_path: Path) -> None:
    docx = Document()
    docx.add_heading("附件", level=1)
    docx.add_paragraph("第一条 住宿服务")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "税率"
    docx_path = tmp_path / "a.docx"
    docx.save(docx_path)
    parsed_docx = parse_asset(docx_path, asset_id="asset-1", asset_sha256=SHA)
    assert parsed_docx.quality_status == "accepted"
    assert {block.kind.value for block in parsed_docx.blocks} >= {"heading", "clause", "table"}

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["项目", "税率"])
    sheet.append(["住宿服务", "6%"])
    xlsx_path = tmp_path / "a.xlsx"
    workbook.save(xlsx_path)
    parsed_xlsx = parse_asset(xlsx_path, asset_id="asset-2", asset_sha256=SHA)
    assert parsed_xlsx.blocks[0].table_id == f"sheet:{sheet.title}"
    assert "住宿服务 | 6%" in parsed_xlsx.blocks[0].text


def test_scanned_pdf_uses_bounded_ocr_and_marks_confidence(tmp_path: Path) -> None:
    document = pymupdf.open()
    document.new_page()
    raw = document.tobytes()
    document.close()
    path = tmp_path / "scan.pdf"
    path.write_bytes(raw)

    parsed = parse_asset(
        path,
        asset_id="asset-scan",
        asset_sha256=SHA,
        ocr_engine=lambda _: ("扫描住宿服务正文", (0.95,)),
    )
    assert parsed.ocr_status is OcrStatus.ACCEPTED
    assert parsed.quality_status == "accepted"
    assert parsed.blocks[0].page_number == 1


def test_native_pdf_parser_preserves_text(tmp_path: Path) -> None:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Hotel accommodation service evidence")
    path = tmp_path / "native.pdf"
    document.save(path)
    document.close()

    parsed = parse_asset(path, asset_id="asset-native-pdf", asset_sha256=SHA)

    assert parsed.ocr_status is OcrStatus.NOT_APPLIED
    assert parsed.quality_status == "accepted"
    assert parsed.blocks[0].page_number == 1
    assert "Hotel accommodation service evidence" in parsed.blocks[0].text


def test_parser_removes_legacy_hyperlink_field_code_without_losing_visible_text(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    class Result:
        text = '旅游服务包括HYPERLINK "http://example.invalid" \\t "_blank"交通、住宿。'

    import legacy_doc

    monkeypatch.setattr(legacy_doc, "extract_text", lambda _: Result())
    path = tmp_path / "legacy.doc"
    path.write_bytes(b"unused by fake")
    parsed = parse_asset(path, asset_id="asset-legacy", asset_sha256=SHA)
    assert parsed.blocks[0].text == "旅游服务包括交通、住宿。"


def test_legacy_parser_preserves_heading_clause_and_table_boundaries(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class Result:
        text = "一、适用范围\r\n第一条 住宿服务。\r\n项目\t税率\r\n住宿服务\t6%"

    import legacy_doc

    monkeypatch.setattr(legacy_doc, "extract_text", lambda _: Result())
    path = tmp_path / "legacy.doc"
    path.write_bytes(b"unused by fake")

    parsed = parse_asset(path, asset_id="asset-legacy", asset_sha256=SHA)

    assert [block.kind.value for block in parsed.blocks] == ["heading", "clause", "table"]
    assert parsed.blocks[1].clause_id == "第一条"
    assert parsed.blocks[1].section_path == ("一、适用范围",)
    assert parsed.blocks[2].table_id == "table-1"
    assert parsed.blocks[2].text == "项目 | 税率\n住宿服务 | 6%"


def test_chunking_is_deterministic_and_bounded(tmp_path: Path) -> None:
    path = tmp_path / "a.html"
    path.write_text("<html><body><p>" + "住宿服务。" * 400 + "</p></body></html>", encoding="utf-8")
    parsed = parse_asset(path, asset_id="asset-1", asset_sha256=SHA)
    first = chunk_document(parsed, document_id="tax-1", asset_version=SHA)
    second = chunk_document(parsed, document_id="tax-1", asset_version=SHA)
    assert first == second
    assert all(1 <= len(chunk.content) <= 1600 for chunk in first)


def test_chunk_contract_rejects_content_hash_mismatch() -> None:
    with pytest.raises(ValidationError, match="content SHA-256 mismatch"):
        CorpusChunk(
            schema_version=1,
            chunk_id="chunk-1",
            document_id="doc-1",
            asset_id="asset-1",
            asset_version=SHA,
            ordinal=1,
            content="正文",
            content_sha256=SHA,
            ocr_applied=False,
            ocr_confidence_status=OcrStatus.NOT_APPLIED,
        )


def test_processing_result_counts_are_closed() -> None:
    with pytest.raises(ValidationError, match="counts must equal"):
        ProcessingResult(
            schema_version=1,
            run_id="run-1",
            asset_count=2,
            accepted_asset_count=1,
            review_required_asset_count=0,
            rejected_asset_count=0,
            chunk_count=1,
        )


def test_processing_rejects_manifest_path_outside_workspace(tmp_path: Path) -> None:
    workspace = tmp_path / "workspace"
    outside = tmp_path / "outside.html"
    raw = b"<html><body><p>outside</p></body></html>"
    outside.write_bytes(raw)
    manifest = AssetManifest(
        schema_version=1,
        asset_id="asset-1",
        asset_version=sha256_bytes(raw),
        parent_document_id="doc-1",
        source_url="https://www.chinatax.gov.cn/a.html",
        source_final_url="https://www.chinatax.gov.cn/a.html",
        fetched_at_utc=datetime.now(UTC),
        filename="a.html",
        source_extension=".html",
        extension=".html",
        detected_mime="text/html",
        sha256=sha256_bytes(raw),
        byte_count=len(raw),
        storage_relative_path="../outside.html",
        official_source_proof="official test fixture",
    )
    manifest_path = tmp_path / "manifest.jsonl"
    write_jsonl(manifest_path, (manifest,))

    with pytest.raises(StateConflict, match="escaped workspace"):
        process_assets(
            asset_manifest_path=manifest_path,
            workspace=workspace,
            run_id="run-1",
            enable_ocr=False,
        )


def test_embedding_contract_uses_runtime_shape() -> None:
    def handler(_: httpx.Request) -> httpx.Response:
        return httpx.Response(200, json={"dim": 1024, "vectors": [[0.0] * 1024]})

    with httpx.Client(transport=httpx.MockTransport(handler)) as client:
        vectors = EmbeddingClient("http://bge.local", client=client).embed(("text",))
    assert len(vectors) == 1 and len(vectors[0]) == 1024


def test_candidate_clone_copies_analysis_without_generated_or_write_block_settings() -> None:
    requests: list[httpx.Request] = []

    def handler(request: httpx.Request) -> httpx.Response:
        requests.append(request)
        if request.method == "GET" and request.url.path == "/source":
            return httpx.Response(200, json={"source": {"settings": {"index": {"uuid": "uuid-1", "provided_name": "source", "blocks": {"write": "true"}, "number_of_shards": "1", "analysis": {"analyzer": {"a": {"type": "standard"}}}}}, "mappings": {"_meta": {"mapping_version": "v1"}, "properties": {"content": {"type": "text"}}}}})
        if request.method == "HEAD":
            return httpx.Response(404)
        if request.method == "PUT" and request.url.path == "/candidate":
            return httpx.Response(200, json={"acknowledged": True})
        if request.method == "POST" and request.url.path == "/_reindex":
            return httpx.Response(200, json={"created": 1, "total": 1, "failures": []})
        raise AssertionError(f"unexpected {request.method} {request.url}")

    with httpx.Client(transport=httpx.MockTransport(handler)) as client:
        ElasticsearchCandidateAdmin("http://es.local", client=client).create_and_clone(
            source_index="source", source_uuid="uuid-1", candidate_index="candidate"
        )
    creation = next(request for request in requests if request.method == "PUT")
    body = json.loads(creation.content)
    assert body["settings"]["analysis"]["analyzer"]["a"]["type"] == "standard"
    assert "blocks" not in body["settings"] and "uuid" not in body["settings"]
    assert body["mappings"]["_meta"]["mapping_version"] == "agent-knowledge-tax-v2-corpus-a1"


def test_index_document_preserves_acl_and_relationship() -> None:
    path = Path(__file__).parent / "fixture.html"
    path.write_text("<html><body><p>住宿服务</p></body></html>", encoding="utf-8")
    try:
        parsed = parse_asset(path, asset_id="asset-1", asset_sha256=SHA)
    finally:
        path.unlink(missing_ok=True)
    chunk = chunk_document(parsed, document_id="tax-parent", asset_version=SHA)[0]
    result = to_index_document(
        chunk=chunk,
        vector=[0.0] * 1024,
        parent_source={"aclRef": "public:tax_policy", "aclVersion": "public-v1", "visibility": "PUBLIC", "channel": "财税文件", "tenantId": "tenant-local", "corpusId": "tax_policy", "domain": "tax_policy", "materialType": "tax_policy"},
        attachment_title="附件一",
        source_url="https://www.chinatax.gov.cn/a.docx",
        source_final_url="https://www.chinatax.gov.cn/a.docx",
        source_fetched_at=datetime.now(UTC).isoformat(),
        parser_version="python-docx-1.2.0",
        chunk_count=1,
    )
    assert result["aclRef"] == "public:tax_policy"
    assert result["parentDocumentId"] == "tax-parent"
    assert result["embedding"] == [0.0] * 1024


def test_alias_switch_checks_exact_preconditions_and_can_rollback() -> None:
    target = {"value": "old"}
    uuids = {"old": "old-uuid", "candidate": "new-uuid"}

    def handler(request: httpx.Request) -> httpx.Response:
        path = request.url.path
        if path == "/_alias/read":
            return httpx.Response(200, json={target["value"]: {"aliases": {"read": {"is_write_index": False}}}})
        if path in {"/old", "/candidate"}:
            index = path[1:]
            return httpx.Response(200, json={index: {"settings": {"index": {"uuid": uuids[index]}}}})
        if path in {"/old/_alias", "/candidate/_alias"}:
            index = path.split("/")[1]
            aliases = {"read": {"is_write_index": False}} if target["value"] == index else {}
            return httpx.Response(200, json={index: {"aliases": aliases}})
        if path == "/_aliases":
            body = json.loads(request.content)
            target["value"] = body["actions"][1]["add"]["index"]
            return httpx.Response(200, json={"acknowledged": True})
        raise AssertionError(path)

    with httpx.Client(transport=httpx.MockTransport(handler)) as client:
        manager = AliasReleaseManager("http://es.local", client=client)
        published = manager.switch(alias="read", expected_from_index="old", expected_from_uuid="old-uuid", target_index="candidate", target_uuid="new-uuid", phase="published")
        rolled_back = manager.switch(alias="read", expected_from_index="candidate", expected_from_uuid="new-uuid", target_index="old", target_uuid="old-uuid", phase="rolled_back")
    assert published.phase == "published"
    assert rolled_back.phase == "rolled_back"
    assert target["value"] == "old"


def test_release_cli_rejects_existing_journal_before_http(tmp_path: Path) -> None:
    output = tmp_path / "runs" / "run-1" / "release-journal.v1.jsonl"
    output.parent.mkdir(parents=True)
    output.write_text("frozen\n", encoding="utf-8")

    with pytest.raises(SystemExit, match="already exists"):
        cli_main(
            [
                "release-rehearsal",
                "--workspace", str(tmp_path),
                "--run-id", "run-1",
                "--es-endpoint", "http://127.0.0.1:1",
                "--alias", "read",
                "--old-index", "old",
                "--old-uuid", "old-uuid",
                "--candidate-index", "candidate",
                "--candidate-uuid", "candidate-uuid",
            ]
        )
    assert output.read_text(encoding="utf-8") == "frozen\n"


def test_uat_result_contract_is_finite_and_consistent() -> None:
    cases = tuple(
        StageAUatCase(
            case_id=f"UAT-KCORPUS-A-{number:02d}",
            evidence_kind="automated_test",
            evidence_refs=("pytest:test",),
            status="passed",
            failure_reason="none",
        )
        for number in range(1, 15)
    )
    result = StageAUatResult(
        schema_version=1,
        run_id="run-1",
        current_alias="read",
        old_index="old",
        candidate_index="candidate",
        candidate_index_uuid="uuid",
        p0_document_count=3,
        p0_attachment_count=4,
        p0_chunk_count=31,
        model_outbound_count=0,
        business_call_count=0,
        cases=cases,
        passed_count=14,
        failed_count=0,
        stage_b_findings=("domain_selection", "query_rewrite", "ranking"),
        conclusion="passed",
        completed_at_utc=datetime.now(UTC),
    )
    assert result.conclusion == "passed"
    with pytest.raises(ValidationError):
        StageAUatResult(**{**result.model_dump(), "passed_count": 13})


def test_committed_uat_schema_is_closed() -> None:
    schema = json.loads((Path(__file__).parents[1] / "schemas" / "stage-a-uat-result-v1.schema.json").read_text(encoding="utf-8"))
    assert schema["additionalProperties"] is False
    assert schema["$defs"]["case"]["additionalProperties"] is False
