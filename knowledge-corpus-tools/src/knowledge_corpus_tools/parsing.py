from __future__ import annotations

import io
import re
import zipfile
from collections.abc import Callable
from pathlib import Path
from typing import Any, Literal, Protocol, cast
from urllib.parse import unquote, urljoin, urlsplit

from bs4 import BeautifulSoup
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import pymupdf
from openpyxl import load_workbook  # type: ignore[import-untyped]
from openpyxl.utils.exceptions import InvalidFileException  # type: ignore[import-untyped]
from pypdf import PdfReader
from pypdf.errors import PdfReadError
import xlrd  # type: ignore[import-untyped]
from xlrd.biffh import XLRDError  # type: ignore[import-untyped]

from .errors import ContractError
from .models import AttachmentReference, BlockKind, OcrStatus, ParsedBlock, ParsedDocument
from .safety import ALLOWED_EXTENSIONS, extension_from_name, normalize_official_url, validate_zip_container

CLAUSE_PATTERN = re.compile(r"^(第[一二三四五六七八九十百千0-9]+条)\s*")
HEADING_PATTERN = re.compile(
    r"^(?:附件(?:[一二三四五六七八九十百千0-9]+)?[:：]?|"
    r"第[一二三四五六七八九十百千0-9]+章(?:\s|$)|"
    r"[一二三四五六七八九十百千]+、)"
)
OcrPage = Callable[[bytes], tuple[str, tuple[float, ...]]]
_pymupdf: Any = pymupdf


class OcrEngine(Protocol):
    def __call__(self, image: bytes) -> tuple[str, tuple[float, ...]]: ...


class RapidOcrEngine:
    """Lazy local OCR adapter; model files are local and no network call is made."""

    def __init__(self) -> None:
        from rapidocr import RapidOCR

        self._engine: Any = RapidOCR()

    def __call__(self, image: bytes) -> tuple[str, tuple[float, ...]]:
        result = self._engine(image)
        texts = tuple(result.txts or ())
        scores = tuple(float(score) for score in (result.scores or ()))
        if len(texts) != len(scores) or any(not 0.0 <= score <= 1.0 for score in scores):
            raise ContractError("invalid OCR output")
        return "\n".join(texts), scores


def _clean(text: str) -> str:
    text = re.sub(r'HYPERLINK\s+"[^"]+"\s+\\t\s+"[^"]+"', "", text)
    return "\n".join(line.strip() for line in text.replace("\r", "\n").split("\n") if line.strip())


def _block(ordinal: int, kind: BlockKind, text: str, **kwargs: Any) -> ParsedBlock:
    return ParsedBlock(ordinal=ordinal, kind=kind, text=_clean(text), **kwargs)


def _parse_structured_text(
    text: str,
    *,
    page_number: int | None = None,
    start_ordinal: int = 1,
) -> list[ParsedBlock]:
    """Preserve line, clause, heading and tabular boundaries from plain text."""

    normalized = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x0c", "\n")
    blocks: list[ParsedBlock] = []
    section_path: tuple[str, ...] = ()
    table_rows: list[str] = []
    table_number = 0

    def append_table() -> None:
        nonlocal table_number
        if not table_rows:
            return
        table_number += 1
        blocks.append(
            _block(
                start_ordinal + len(blocks),
                BlockKind.TABLE,
                "\n".join(table_rows),
                section_path=section_path,
                page_number=page_number,
                table_id=f"table-{table_number}",
            )
        )
        table_rows.clear()

    for raw_line in normalized.split("\n"):
        if "\t" in raw_line:
            cells = tuple(cell.strip() for cell in raw_line.split("\t"))
            if any(cells):
                table_rows.append(" | ".join(cells))
            continue
        append_table()
        line = _clean(raw_line)
        if not line:
            continue
        clause_match = CLAUSE_PATTERN.match(line)
        is_heading = clause_match is None and len(line) <= 120 and HEADING_PATTERN.match(line) is not None
        if is_heading:
            section_path = (line,)
            kind = BlockKind.HEADING
        elif clause_match is not None:
            kind = BlockKind.CLAUSE
        else:
            kind = BlockKind.PARAGRAPH
        blocks.append(
            _block(
                start_ordinal + len(blocks),
                kind,
                line,
                section_path=section_path,
                page_number=page_number,
                clause_id=clause_match.group(1) if clause_match else None,
            )
        )
    append_table()
    return blocks


def parse_asset(
    path: Path,
    *,
    asset_id: str,
    asset_sha256: str,
    ocr_engine: OcrEngine | None = None,
) -> ParsedDocument:
    suffix = path.suffix.lower()
    raw = path.read_bytes()
    try:
        if suffix in {".html", ".htm"}:
            blocks = _parse_html(raw)
            parser_name = "beautifulsoup4-html"
            parser_version = "4.15.0"
        elif suffix == ".pdf":
            blocks = _parse_pdf(raw)
            parser_name = "pypdf-structured"
            parser_version = "6.16.2+structure-v1"
            ocr_status = OcrStatus.NOT_APPLIED
            if not blocks and ocr_engine is not None:
                blocks, ocr_status = _parse_scanned_pdf(raw, ocr_engine)
                parser_name = "pymupdf+rapidocr"
                parser_version = "1.28.2+3.9.2"
        elif suffix == ".docx":
            with zipfile.ZipFile(io.BytesIO(raw)) as archive:
                validate_zip_container(archive)
            blocks = _parse_docx(raw)
            parser_name = "python-docx"
            parser_version = "1.2.0"
        elif suffix == ".xlsx":
            with zipfile.ZipFile(io.BytesIO(raw)) as archive:
                validate_zip_container(archive)
            blocks = _parse_xlsx(raw)
            parser_name = "openpyxl"
            parser_version = "3.1.5"
        elif suffix == ".xls":
            blocks = _parse_xls(raw)
            parser_name = "xlrd"
            parser_version = "2.0.2"
        elif suffix == ".doc":
            blocks = _parse_legacy_doc(path)
            parser_name = "legacy-doc-structured"
            parser_version = "0.2.1+structure-v1"
        else:
            raise ContractError(f"unsupported parser extension: {suffix}")
    except (zipfile.BadZipFile, PdfReadError, PackageNotFoundError, InvalidFileException, XLRDError) as exc:
        raise ContractError("asset parser rejected malformed content") from exc
    if suffix != ".pdf":
        ocr_status = OcrStatus.NOT_APPLIED
    quality = "accepted" if blocks and ocr_status is not OcrStatus.REJECTED else "rejected"
    if ocr_status is OcrStatus.REVIEW_REQUIRED:
        quality = "review_required"
    reasons = () if quality == "accepted" else (("ocr_review_required",) if quality == "review_required" else ("empty_parsed_body",))
    return ParsedDocument(
        schema_version=1,
        asset_id=asset_id,
        asset_sha256=asset_sha256,
        parser_name=parser_name,
        parser_version=parser_version,
        ocr_status=ocr_status,
        blocks=tuple(blocks),
        quality_status=cast(Literal["accepted", "review_required", "rejected"], quality),
        quality_reasons=reasons,
    )


def _parse_html(raw: bytes) -> list[ParsedBlock]:
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    root = soup.select_one(".TRS_Editor, .article-content, #zoom, .content") or soup.body or soup
    blocks: list[ParsedBlock] = []
    for element in root.find_all(["h1", "h2", "h3", "h4", "p", "table"], recursive=True):
        text = element.get_text(" | " if element.name == "table" else " ", strip=True)
        if not text:
            continue
        kind = BlockKind.TABLE if element.name == "table" else BlockKind.HEADING if element.name.startswith("h") else BlockKind.CLAUSE if CLAUSE_PATTERN.match(text) else BlockKind.PARAGRAPH
        kwargs: dict[str, Any] = {}
        if kind is BlockKind.TABLE:
            kwargs["table_id"] = f"table-{len(blocks) + 1}"
        match = CLAUSE_PATTERN.match(text)
        if match:
            kwargs["clause_id"] = match.group(1)
        blocks.append(_block(len(blocks) + 1, kind, text, **kwargs))
    return blocks


def _parse_pdf(raw: bytes) -> list[ParsedBlock]:
    reader = PdfReader(io.BytesIO(raw), strict=True)
    blocks: list[ParsedBlock] = []
    for page_number, page in enumerate(reader.pages, start=1):
        blocks.extend(
            _parse_structured_text(
                page.extract_text() or "",
                page_number=page_number,
                start_ordinal=len(blocks) + 1,
            )
        )
    return blocks


def _parse_scanned_pdf(raw: bytes, engine: OcrEngine) -> tuple[list[ParsedBlock], OcrStatus]:
    try:
        document: Any = _pymupdf.open(stream=raw, filetype="pdf")
    except Exception as exc:
        raise ContractError("invalid PDF for OCR") from exc
    if document.page_count > 200:
        document.close()
        raise ContractError("OCR page limit exceeded")
    blocks: list[ParsedBlock] = []
    scores: list[float] = []
    try:
        for page_number, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=_pymupdf.Matrix(2.0, 2.0), alpha=False)
            image = pixmap.tobytes("png")
            text, page_scores = engine(image)
            scores.extend(page_scores)
            blocks.extend(
                _parse_structured_text(
                    text,
                    page_number=page_number,
                    start_ordinal=len(blocks) + 1,
                )
            )
    finally:
        document.close()
    if not blocks or not scores:
        return blocks, OcrStatus.REJECTED
    minimum = min(scores)
    average = sum(scores) / len(scores)
    if minimum < 0.50 or average < 0.75:
        return blocks, OcrStatus.REVIEW_REQUIRED
    return blocks, OcrStatus.ACCEPTED


def discover_attachments(
    raw: bytes,
    *,
    source_page_url: str,
    parent_document_id: str,
    allowed_hosts: frozenset[str],
) -> tuple[AttachmentReference, ...]:
    soup = BeautifulSoup(raw, "html.parser")
    root = soup.select_one(".TRS_Editor, .article-content, #zoom, .content") or soup.body or soup
    references: list[AttachmentReference] = []
    seen: set[str] = set()
    for anchor in root.find_all("a", href=True):
        href = str(anchor.get("href", "")).strip()
        if not href:
            continue
        joined = urljoin(source_page_url, href)
        path_name = Path(unquote(urlsplit(joined).path)).name
        label = anchor.get_text(" ", strip=True)
        filename = path_name or label
        try:
            extension = extension_from_name(filename)
        except Exception:
            continue
        if extension not in ALLOWED_EXTENSIONS - {".html", ".htm"}:
            continue
        candidate = normalize_official_url(joined, allowed_hosts)
        if candidate in seen:
            continue
        if label:
            filename = f"{label}{extension}" if not label.lower().endswith(extension) else label
        seen.add(candidate)
        references.append(
            AttachmentReference(
                schema_version=1,
                parent_document_id=parent_document_id,
                source_page_url=source_page_url,
                attachment_url=candidate,
                filename=filename,
                extension=cast(Literal[".pdf", ".doc", ".docx", ".xls", ".xlsx"], extension),
                ordinal=len(references) + 1,
            )
        )
    return tuple(references)


def _parse_docx(raw: bytes) -> list[ParsedBlock]:
    document = Document(io.BytesIO(raw))
    blocks: list[ParsedBlock] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").lower()
        match = CLAUSE_PATTERN.match(text)
        kind = BlockKind.HEADING if "heading" in style or "标题" in style else BlockKind.CLAUSE if match else BlockKind.PARAGRAPH
        blocks.append(_block(len(blocks) + 1, kind, text, clause_id=match.group(1) if match else None))
    for table_number, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        text = "\n".join(row for row in rows if row.strip(" |"))
        if text:
            blocks.append(_block(len(blocks) + 1, BlockKind.TABLE, text, table_id=f"table-{table_number}"))
    return blocks


def _parse_xlsx(raw: bytes) -> list[ParsedBlock]:
    workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=False, keep_links=False)
    blocks: list[ParsedBlock] = []
    for worksheet in workbook.worksheets:
        rows: list[str] = []
        for row in worksheet.iter_rows(values_only=False):
            values = ["" if cell.value is None else str(cell.value) for cell in row]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            blocks.append(_block(len(blocks) + 1, BlockKind.TABLE, "\n".join(rows), table_id=f"sheet:{worksheet.title}"))
    workbook.close()
    return blocks


def _parse_xls(raw: bytes) -> list[ParsedBlock]:
    workbook = xlrd.open_workbook(file_contents=raw, on_demand=True)
    blocks: list[ParsedBlock] = []
    for sheet in workbook.sheets():
        rows = [" | ".join(str(sheet.cell_value(r, c)) for c in range(sheet.ncols)) for r in range(sheet.nrows)]
        text = "\n".join(row for row in rows if row.strip(" |"))
        if text:
            blocks.append(_block(len(blocks) + 1, BlockKind.TABLE, text, table_id=f"sheet:{sheet.name}"))
    workbook.release_resources()
    return blocks


def _parse_legacy_doc(path: Path) -> list[ParsedBlock]:
    try:
        from legacy_doc import extract_text  # type: ignore[import-untyped]
    except (ImportError, AttributeError) as exc:
        raise ContractError("legacy DOC parser unavailable") from exc
    result = extract_text(path.read_bytes())
    return _parse_structured_text(str(result.text))
